"""Template analysis tool — identify placeholders and extract formatting context.

Supports two backends:
- python-docx: Full structured parsing of paragraphs, tables, headers/footers.
- XML fallback: Regex-based scanning of raw document XML when python-docx is unavailable.
"""

from __future__ import annotations

import logging
import re
import zipfile
from pathlib import Path

from agent_doc_filler.models import (
    DocumentStats,
    HeadingInfo,
    PlaceholderInfo,
    SectionContent,
    TableInfo,
    TemplateAnalysis,
)

logger = logging.getLogger(__name__)

# Placeholder pattern: {{name}}
PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")

# Heading style pattern: matches "Heading 1", "Heading 2", etc.
HEADING_STYLE_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)

# Chinese heading prefix pattern: matches "一、", "（一）", "1.", "1.1", etc.
HEADING_PREFIX_RE = re.compile(
    r"^(?:"
    r"[一二三四五六七八九十]+[、.]"  # 一、二、
    r"|[（(][一二三四五六七八九十]+[）)]"  # （一）
    r"|\d+[.\s]"  # 1. or 1<space>
    r")"
)

MAX_PREVIEW_LENGTH = 200
MAX_TABLE_PREVIEW_ROWS = 3


def _extract_formatting_from_run(run: object) -> dict:
    """Extract formatting properties from a python-docx Run object.

    Returns a dict of non-None formatting attributes, or empty dict if
    python-docx is not available.
    """
    fmt: dict = {}
    try:
        if run.font.name:
            fmt["font_name"] = run.font.name
        if run.font.size:
            fmt["font_size"] = run.font.size.pt
        if run.font.bold is not None:
            fmt["bold"] = run.font.bold
        if run.font.italic is not None:
            fmt["italic"] = run.font.italic
        if run.font.underline is not None:
            fmt["underline"] = run.font.underline
        if run.font.color and run.font.color.rgb:
            fmt["font_color"] = str(run.font.color.rgb)
    except Exception:
        logger.exception("Failed to extract formatting from docx Run object")
    return fmt


def _guess_field_type(name: str) -> str:
    """Infer a field type hint from the placeholder name."""
    name_lower = name.lower()
    if "date" in name_lower or "time" in name_lower:
        return "date"
    if any(kw in name_lower for kw in ("amount", "price", "total", "count", "number", "num", "qty")):
        return "number"
    if "image" in name_lower or "photo" in name_lower or "logo" in name_lower:
        return "image_ref"
    return "text"


def _get_heading_level(paragraph) -> int | None:
    """Determine if a paragraph is a heading and return its level (1-9), or None."""
    style = paragraph.style
    if style is None:
        return None

    # Check named style (e.g. "Heading 1", "标题 1")
    style_name = style.name or ""
    match = HEADING_STYLE_RE.match(style_name)
    if match:
        return int(match.group(1))

    # Chinese Word heading styles
    cn_match = re.match(r"^标题\s*(\d+)$", style_name)
    if cn_match:
        return int(cn_match.group(1))

    # Bold text at start of document sections can be a visual heading
    # Only treat as heading if font is significantly larger or bold
    if style_name == "Normal" and paragraph.runs:
        text = paragraph.text.strip()
        if text and HEADING_PREFIX_RE.match(text):
            return 2  # Treat Chinese numbered items as level 2

    return None


def _count_images(doc) -> int:
    """Count inline images in the document by scanning relationships."""
    try:
        count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                count += 1
        return count
    except Exception:
        return 0


def _count_words(text: str) -> int:
    """Count words in mixed Chinese/English text."""
    if not text:
        return 0
    # Chinese characters count as individual words
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    # Remove Chinese chars, then count English words
    remaining = re.sub(r"[\u4e00-\u9fff]", " ", text)
    english_words = len(remaining.split())
    return chinese_chars + english_words


def _analyze_with_docx(template_path: str) -> TemplateAnalysis:
    """Full analysis using python-docx library."""
    from docx import Document

    doc = Document(template_path)
    placeholders: list[PlaceholderInfo] = []
    seen_names: set[str] = set()
    headings: list[HeadingInfo] = []
    sections: list[SectionContent] = []
    tables: list[TableInfo] = []
    total_chars = 0
    total_words = 0
    image_count = _count_images(doc)

    def _scan_text(text: str, formatting: dict | None = None) -> None:
        """Scan a text string for placeholders and add them."""
        for match in PLACEHOLDER_RE.finditer(text):
            name = match.group(1)
            if name in seen_names:
                continue
            seen_names.add(name)
            placeholders.append(
                PlaceholderInfo(
                    name=name,
                    field_type=_guess_field_type(name),
                    description="",
                    required=True,
                    default=None,
                    formatting=formatting,
                )
            )

    # --- Build headings and sections ---
    current_section_text: list[str] = []
    current_heading = ""
    current_level = 0
    current_para_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        total_chars += len(text)
        total_words += _count_words(text)

        # Scan for placeholders
        for run in para.runs:
            if PLACEHOLDER_RE.search(run.text):
                _scan_text(run.text, _extract_formatting_from_run(run))

        # Detect heading
        level = _get_heading_level(para)
        if level is not None and text:
            # Flush previous section
            if current_heading or current_para_count > 0:
                preview = "\n".join(current_section_text)[:MAX_PREVIEW_LENGTH]
                sections.append(
                    SectionContent(
                        heading=current_heading,
                        level=current_level,
                        paragraph_count=current_para_count,
                        preview=preview,
                    )
                )
            # Start new section
            headings.append(HeadingInfo(level=level, text=text))
            current_heading = text
            current_level = level
            current_section_text = []
            current_para_count = 0
        elif text:
            current_para_count += 1
            if len(current_section_text) < 5:
                current_section_text.append(text)

    # Flush last section
    if current_heading or current_para_count > 0:
        preview = "\n".join(current_section_text)[:MAX_PREVIEW_LENGTH]
        sections.append(
            SectionContent(
                heading=current_heading,
                level=current_level,
                paragraph_count=current_para_count,
                preview=preview,
            )
        )

    # --- Scan tables ---
    for idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_texts)
            # Scan cells for placeholders
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if PLACEHOLDER_RE.search(run.text):
                            _scan_text(run.text, _extract_formatting_from_run(run))

        header_row = rows_data[0] if rows_data else []
        num_cols = max((len(r) for r in rows_data), default=0)
        preview = rows_data[:MAX_TABLE_PREVIEW_ROWS + 1]  # header + preview rows

        tables.append(
            TableInfo(
                index=idx,
                rows=len(rows_data),
                cols=num_cols,
                header_row=header_row,
                preview=preview,
            )
        )

    # Scan headers and footers for placeholders
    for section in doc.sections:
        for para in section.header.paragraphs:
            _scan_text(para.text)
        for para in section.footer.paragraphs:
            _scan_text(para.text)

    # --- Document-level style info ---
    style_info: dict = {}
    if doc.styles:
        try:
            normal = doc.styles["Normal"]
            if normal.font.name:
                style_info["default_font"] = normal.font.name
            if normal.font.size:
                style_info["default_size"] = normal.font.size.pt
        except (KeyError, AttributeError):
            pass

    # Core properties
    metadata: dict = {}
    metadata["section_count"] = len(doc.sections)
    core_props = doc.core_properties
    if core_props.title:
        metadata["title"] = core_props.title
    if core_props.author:
        metadata["author"] = core_props.author
    if core_props.created:
        metadata["created"] = str(core_props.created)
    if core_props.modified:
        metadata["modified"] = str(core_props.modified)
    if core_props.subject:
        metadata["subject"] = core_props.subject

    stats = DocumentStats(
        total_paragraphs=len(doc.paragraphs),
        total_tables=len(doc.tables),
        total_images=image_count,
        total_characters=total_chars,
        total_words=total_words,
        heading_count=len(headings),
    )

    return TemplateAnalysis(
        template_path=template_path,
        placeholders=placeholders,
        headings=headings,
        sections=sections,
        tables=tables,
        stats=stats,
        style_info=style_info,
        metadata=metadata,
    )


def _analyze_xml_fallback(template_path: str) -> TemplateAnalysis:
    """Fallback analysis using XML regex when python-docx is unavailable.

    Reads the raw XML inside the .docx zip and scans for {{placeholder}} patterns.
    Also extracts basic structure information from the XML.
    """
    placeholders: list[PlaceholderInfo] = []
    seen_names: set[str] = set()
    headings: list[HeadingInfo] = []
    total_chars = 0

    with zipfile.ZipFile(template_path, "r") as zf:
        # Scan main document
        xml_files = [n for n in zf.namelist() if n.endswith(".xml")]
        image_count = sum(
            1 for n in zf.namelist()
            if n.startswith("word/media/")
        )
        for xml_name in xml_files:
            content = zf.read(xml_name).decode("utf-8")
            # Remove XML tags to get raw text for placeholder scanning
            text_only = re.sub(r"<[^>]+>", " ", content)
            total_chars += len(text_only.strip())

            for match in PLACEHOLDER_RE.finditer(text_only):
                name = match.group(1)
                if name in seen_names:
                    continue
                seen_names.add(name)
                placeholders.append(
                    PlaceholderInfo(
                        name=name,
                        field_type=_guess_field_type(name),
                        description="",
                        required=True,
                        default=None,
                        formatting=None,
                    )
                )

        # Try to extract headings from document.xml
        doc_xml_name = "word/document.xml"
        if doc_xml_name in zf.namelist():
            doc_xml = zf.read(doc_xml_name).decode("utf-8")
            # Look for text in heading-style paragraphs
            # This is a rough heuristic — look for <w:pStyle w:val="HeadingN"/>
            heading_pattern = re.compile(
                r'<w:pStyle\s+w:val="Heading(\d+)"[^/]*/>'
                r'.*?<w:t[^>]*>([^<]+)</w:t>',
                re.DOTALL,
            )
            for match in heading_pattern.finditer(doc_xml):
                level = int(match.group(1))
                text = match.group(2).strip()
                if text:
                    headings.append(HeadingInfo(level=level, text=text))

    return TemplateAnalysis(
        template_path=template_path,
        placeholders=placeholders,
        headings=headings,
        stats=DocumentStats(
            total_characters=total_chars,
            total_images=image_count,
            heading_count=len(headings),
        ),
        style_info={},
        metadata={"backend": "xml_fallback"},
    )


def analyze_template(template_path: str) -> TemplateAnalysis:
    """Analyze a .docx template file to identify all placeholders.

    Uses python-docx if available for rich structural analysis (paragraphs,
    tables, headers/footers with formatting context). Falls back to XML regex
    scanning when python-docx is not installed.

    Args:
        template_path: Path to the .docx template file.

    Returns:
        TemplateAnalysis with all discovered placeholders and style info.

    Raises:
        FileNotFoundError: If template_path does not exist.
        ValueError: If the file is not a .docx file.
    """
    path = Path(template_path)
    if not path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")
    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Expected .docx file, got: {path.suffix}")

    try:
        return _analyze_with_docx(template_path)
    except ImportError:
        return _analyze_xml_fallback(template_path)
    except Exception:
        # python-docx may fail on minimal/malformed docx files
        # (missing [Content_Types].xml, corrupt structure, etc.)
        logger.debug("python-docx failed, falling back to XML", exc_info=True)
        return _analyze_xml_fallback(template_path)
