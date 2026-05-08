"""Template filling tool — replace placeholders with values, preserving styles.

Supports two backends:
- python-docx: Full structured filling with run-level format preservation.
- XML fallback: Simple regex replacement in raw document XML.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from agent_doc_filler.models import FillResult

PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")


def _default_output_path(template_path: str) -> str:
    """Generate default output path with _filled suffix."""
    path = Path(template_path)
    return str(path.with_stem(path.stem + "_filled"))


class _ReplaceState:
    """Mutable state holder for placeholder replacement tracking."""

    __slots__ = ("filled_count", "seen_placeholders")

    def __init__(self) -> None:
        self.filled_count = 0
        self.seen_placeholders: set[str] = set()


def _replace_in_paragraph(paragraph: object, values: dict[str, str], state: _ReplaceState) -> None:
    """Replace placeholders in a single paragraph, preserving run formatting."""
    for run in paragraph.runs:
        text = run.text
        if not PLACEHOLDER_RE.search(text):
            continue

        def _replacer(match: re.Match) -> str:
            name = match.group(1)
            state.seen_placeholders.add(name)
            if name in values:
                state.filled_count += 1
                return values[name]
            return match.group(0)

        run.text = PLACEHOLDER_RE.sub(_replacer, text)


def _replace_in_paragraphs(paragraphs, values: dict[str, str], state: _ReplaceState) -> None:
    """Apply placeholder replacement to an iterable of paragraphs."""
    for para in paragraphs:
        _replace_in_paragraph(para, values, state)


def _replace_placeholders_in_doc(
    doc,
    values: dict[str, str],
) -> tuple[int, set[str]]:
    """Replace all placeholders in document regions. Returns (filled_count, seen_names)."""
    state = _ReplaceState()

    _replace_in_paragraphs(doc.paragraphs, values, state)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, values, state)
    for section in doc.sections:
        _replace_in_paragraphs(section.header.paragraphs, values, state)
        _replace_in_paragraphs(section.footer.paragraphs, values, state)

    return state.filled_count, state.seen_placeholders


def _collect_unfilled(
    doc, seen_placeholders: set[str], values: dict[str, str]
) -> tuple[list[str], list[str]]:
    """Determine unfilled placeholders and detect multi-run placeholders."""
    unfilled: list[str] = [name for name in seen_placeholders if name not in values]
    warnings: list[str] = []

    all_text = " ".join(para.text for para in doc.paragraphs)
    remaining = PLACEHOLDER_RE.search(all_text)
    if remaining:
        name = remaining.group(1)
        if name not in seen_placeholders and name not in values:
            unfilled.append(name)
            warnings.append(
                f"Placeholder '{name}' spans multiple runs and was not filled. "
                "Consider using a single run for this placeholder."
            )

    if unfilled:
        warnings.append(f"Unfilled placeholders: {', '.join(unfilled)}")

    return unfilled, warnings


def _fill_with_docx(
    template_path: str,
    values: dict[str, str],
    output_path: str,
) -> FillResult:
    """Fill template using python-docx with full style preservation."""
    from docx import Document

    doc = Document(template_path)
    filled_count, seen_placeholders = _replace_placeholders_in_doc(doc, values)
    unfilled, warnings = _collect_unfilled(doc, seen_placeholders, values)
    doc.save(output_path)

    return FillResult(
        success=True,
        output_path=output_path,
        filled_count=filled_count,
        unfilled=unfilled,
        warnings=warnings,
    )


def _fill_xml_fallback(
    template_path: str,
    values: dict[str, str],
    output_path: str,
) -> FillResult:
    """Fill template using XML regex replacement (fallback)."""
    filled_count = 0
    unfilled: list[str] = []
    warnings: list[str] = ["Using XML fallback — style preservation may be limited"]

    # Copy the original .docx to output path
    shutil.copy2(template_path, output_path)

    # Read and replace in all XML files within the zip
    with (
        zipfile.ZipFile(output_path, "r") as zin,
        zipfile.ZipFile(output_path + ".tmp", "w") as zout,
    ):
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml"):
                text = data.decode("utf-8")

                # Find all placeholder names in this XML
                found = PLACEHOLDER_RE.findall(text)
                for name in found:
                    if name in values:
                        text = text.replace("{{" + name + "}}", values[name])
                        filled_count += 1
                    else:
                        if name not in unfilled:
                            unfilled.append(name)

                data = text.encode("utf-8")
            zout.writestr(item, data)

    # Replace original with the new file
    Path(output_path + ".tmp").rename(output_path)

    if unfilled:
        warnings.append(f"Unfilled placeholders: {', '.join(unfilled)}")

    return FillResult(
        success=True,
        output_path=output_path,
        filled_count=filled_count,
        unfilled=unfilled,
        warnings=warnings,
    )


def fill_template(
    template_path: str,
    values: dict[str, str],
    output_path: str | None = None,
) -> FillResult:
    """Fill a .docx template with provided values, preserving styles.

    Args:
        template_path: Path to the .docx template file.
        values: Mapping of placeholder names to replacement values.
        output_path: Where to save the filled document. If None, uses
            template_path with a "_filled" suffix.

    Returns:
        FillResult with fill statistics and any warnings.

    Raises:
        FileNotFoundError: If template_path does not exist.
        ValueError: If the file is not a .docx file.
    """
    path = Path(template_path)
    if not path.exists():
        return FillResult(
            success=False,
            output_path=output_path or template_path,
            warnings=[f"Template file not found: {template_path}"],
        )

    if path.suffix.lower() not in (".docx", ".doc"):
        return FillResult(
            success=False,
            output_path=output_path or template_path,
            warnings=[f"Expected .docx file, got: {path.suffix}"],
        )

    if output_path is None:
        output_path = _default_output_path(template_path)

    try:
        return _fill_with_docx(template_path, values, output_path)
    except ImportError:
        return _fill_xml_fallback(template_path, values, output_path)
    except Exception:
        # python-docx may fail on minimal/malformed docx files
        import logging

        logging.getLogger(__name__).debug(
            "python-docx fill failed, falling back to XML", exc_info=True
        )
        return _fill_xml_fallback(template_path, values, output_path)
